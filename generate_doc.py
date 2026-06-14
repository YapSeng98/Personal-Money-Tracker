from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper colours ────────────────────────────────────────────
C_INK    = RGBColor(0x0D, 0x11, 0x17)
C_JADE   = RGBColor(0x00, 0xC8, 0x96)
C_SKY    = RGBColor(0x3B, 0x82, 0xF6)
C_VIOLET = RGBColor(0x8B, 0x5C, 0xF6)
C_CORAL  = RGBColor(0xFF, 0x5B, 0x5B)
C_AMBER  = RGBColor(0xF5, 0x9E, 0x0B)
C_MUTED  = RGBColor(0x6B, 0x72, 0x80)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    tcPr.append(shd)

def heading(text, level=1, color=C_INK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True
    return p

def body(text, bold=False, color=None, size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    return p

def table_with_header(headers, rows, header_bg='00C896'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = C_WHITE
    # data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if (ri % 2) == 1:
                set_cell_bg(cell, 'F7F9FC')
    return t

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Personal Money Tracker')
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = C_INK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Design · Flow · Data Structure · Code Review')
run2.font.size = Pt(13)
run2.font.color.rgb = C_MUTED

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('yapseng98.github.io/Personal-Money-Tracker  |  June 2026')
run3.font.size = Pt(10)
run3.font.color.rgb = C_SKY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. DESIGN STRUCTURE
# ══════════════════════════════════════════════════════════════
heading('1. Design Structure', 1, C_JADE)

heading('1.1 Color Palette (CSS Variables)', 2, C_INK)
body('Defined in :root — all components reference these tokens.', size=Pt(10))
spacer()
table_with_header(
    ['Variable', 'Hex Value', 'Use'],
    [
        ['--ink',     '#0D1117', 'Primary dark text / hero background'],
        ['--ink2',    '#1C2333', 'Secondary dark'],
        ['--ink3',    '#2D3748', 'Tertiary dark'],
        ['--muted',   '#6B7280', 'Secondary text, muted elements'],
        ['--border',  '#E2E8F0', 'Card & input borders'],
        ['--surface', '#F7F9FC', 'Light panel/card backgrounds'],
        ['--white',   '#FFFFFF', 'Pure white'],
        ['--jade',    '#00C896', 'Primary accent — income, success, CTA'],
        ['--jade-d',  '#00A37A', 'Dark jade variant'],
        ['--jade-s',  '#E6FBF5', 'Soft jade panel background'],
        ['--coral',   '#FF5B5B', 'Negative / delete / expense'],
        ['--coral-s', '#FFF0F0', 'Soft coral panel background'],
        ['--amber',   '#F59E0B', 'Warning / edit actions'],
        ['--amber-s', '#FFFBEB', 'Soft amber background'],
        ['--sky',     '#3B82F6', 'Info / links / transfer'],
        ['--sky-s',   '#EFF6FF', 'Soft sky background'],
        ['--violet',  '#8B5CF6', 'Secondary accent / AI badge'],
        ['--violet-s','#F5F3FF', 'Soft violet background'],
    ]
)

spacer()
heading('1.2 Typography', 2)
table_with_header(
    ['Variable', 'Font', 'Use'],
    [
        ['--ff', 'Inter',           'Body text, labels, navigation'],
        ['--fd', 'Space Grotesk',   'Display headings, balance values, KPI numbers'],
        ['--fm', 'JetBrains Mono',  'Currency amounts, monospace values'],
    ]
)

spacer()
heading('1.3 Layout Shell', 2)
code_block(
    '.app-shell (flex row)\n'
    '├── .sidebar (fixed 240px, dark ink background)\n'
    '│   ├── .sidebar-logo (18px display font)\n'
    '│   ├── .nav-section + .nav-item  (×7 views)\n'
    '│   └── .sn-status (ServiceNow connection indicator)\n'
    '└── .main (flex:1)\n'
    '    ├── .topbar (sticky 58px, backdrop blur, z-index:50)\n'
    '    └── .content (scrollable, 28px padding)\n'
    '        └── .view (display:none) → .view.active (visible)'
)

spacer()
heading('1.4 Responsive Breakpoints', 2)
table_with_header(
    ['Breakpoint', 'Changes'],
    [
        ['≤ 900px (mobile)', 'Sidebar slides in as overlay, hamburger button appears, grids collapse to 1 column, .balance-grid stacks'],
        ['≤ 600px',          'KPI strip becomes 2 columns, topbar 14px padding, buttons smaller, CSV shows ⬇ icon'],
        ['≤ 700px height',   'Login PIN pad shrinks to 64px buttons, reduced spacing'],
    ]
)

spacer()
heading('1.5 Key Component Classes', 2)
table_with_header(
    ['Class', 'Description'],
    [
        ['.balance-hero',    'Gradient dark card, radial glow blob, net balance display'],
        ['.kpi-strip',       '4-column grid of KPI cards (transactions, avg spend, savings %, budgets over)'],
        ['.txn-row',         'Flex row: icon + body + amount + actions (opacity:0, visible on hover/touch)'],
        ['.txn-actions',     'Edit/delete icon buttons; @media(hover:none) keeps them always visible on mobile'],
        ['.goal-card',       'White bordered card with .goal-track (gray bar) + .goal-fill (colored gradient)'],
        ['.budget-item',     'Budget card with .budget-track + .budget-fill, alert color when near limit'],
        ['.modal-backdrop',  'Fixed overlay + .modal with fadeUp animation; .open class shows it'],
        ['.btn-jade/coral/amber/sky/ghost', 'Semantic button variants by action type'],
        ['.form-input',      'Unified input style with focus border (sky), 10px border-radius'],
        ['.amount-input-wrap', 'Relative wrapper for .amount-prefix (currency symbol) + number input'],
        ['.curr-sym',        'Currency symbol span — updated dynamically when currency changes'],
        ['.ai-panel',        'Dark gradient panel for Groq AI output with loading animation'],
        ['.pin-dot',         '15px circle; .filled = jade + glow; .error = coral + shake animation'],
    ]
)

spacer()
heading('1.6 Animations', 2)
table_with_header(
    ['Animation', 'Duration', 'Effect'],
    [
        ['fadeUp',   '0.2s ease', 'Modal entrance: opacity 0→1, translateY 16px→0'],
        ['pulse',    '2s infinite', 'SN status dot opacity pulse'],
        ['lbf',      '7–12s ease-in-out', 'Login background blob float (translate + scale)'],
        ['bounce',   '0.9s infinite', 'AI loading dots staggered bounce'],
        ['pinShake', '0.35s',    'Wrong PIN horizontal shake ±10px'],
        ['iconGlow', '3s infinite', 'Login icon box-shadow glow expand/contract'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. FLOW STRUCTURE
# ══════════════════════════════════════════════════════════════
heading('2. Flow Structure', 1, C_JADE)

heading('2.1 App Startup Sequence', 2)
code_block(
    'init()  [runs on page load]\n'
    '├── loadState()          → restore from localStorage[pfmt_state_v2]\n'
    '├── if no saved state    → loadSampleData()  (12 txns, 5 budgets, 3 goals)\n'
    '├── populate settings inputs (SN instance, username, Groq key, currency, language)\n'
    '├── applyLanguage()      → translate all static DOM text\n'
    '├── updateMonthUI()      → set month label and Today button\n'
    '└── show login overlay\n'
    '    ├── no PIN stored    → "setup" mode (first time)\n'
    '    └── PIN exists       → "unlock" mode (returning user)'
)

spacer()
heading('2.2 Navigation Flow', 2)
code_block(
    'showView(name, el)\n'
    '├── hide all .view elements\n'
    '├── remove .active from all .nav-item\n'
    '├── show  #view-{name}, add .active to clicked nav item\n'
    '├── update topbar title from titleKeys map\n'
    '├── show/hide month nav bar (dashboard, budgets, transactions only)\n'
    '├── if name === "transactions" → populateCategoryFilter()\n'
    '└── closeSidebar() + renderAll()'
)

spacer()
heading('2.3 Transaction Modal Flow', 2)
code_block(
    'openTxnModal(id?)\n'
    '├── populate account dropdown from state.accounts\n'
    '│   └── if empty → show "No accounts — add one in Accounts tab"\n'
    '├── if id provided (EDIT mode)\n'
    '│   ├── find txn by id (variable: "txn", NOT "t" — avoid shadowing t())\n'
    '│   ├── fill amount, description, date, notes, category, account\n'
    '│   └── set button text "Update Transaction" (amber)\n'
    '└── if no id (NEW mode)\n'
    '    ├── clear all fields, set date to today\n'
    '    └── set button text "Save Transaction" (jade)\n'
    '\n'
    'saveTransaction()\n'
    '├── validate: amount > 0, description not empty, date present\n'
    '├── if edit → update existing in state.transactions[]\n'
    '├── if new  → prepend { id: nextId++, type, amount, desc, cat, account, date, notes }\n'
    '├── closeTxnModal() + saveState() + renderAll()\n'
    '└── showToast("Transaction saved")'
)

spacer()
heading('2.4 Budget Modal Flow', 2)
code_block(
    'openBudgetModal(id?)\n'
    '├── if edit → lock category dropdown, load spent from getBudgetSpent()\n'
    '└── if new  → allow category selection\n'
    '\n'
    'saveBudget()\n'
    '├── validate: amount > 0\n'
    '├── check no duplicate category (new budgets only)\n'
    '└── create/update → renderAll()'
)

spacer()
heading('2.5 Goal Modal Flow', 2)
code_block(
    'openGoalModal(id?)\n'
    '├── set default target date = today + 1 year\n'
    '├── if edit → show Quick Contribution buttons (+$50, +$100, +$200, +$500)\n'
    '└── if new  → hide Quick Contribution section\n'
    '\n'
    'quickContrib(amount) → adds to goal-current input, shows toast\n'
    '\n'
    'saveGoal()\n'
    '├── validate: name not empty, target > 0\n'
    '└── create/update goal → renderAll()'
)

spacer()
heading('2.6 Account Modal Flow', 2)
code_block(
    'openAccountModal(id?)\n'
    '├── populateInstitutionDropdown(selected)\n'
    '│   ├── build options from state.institutions[]\n'
    '│   └── append "➕ Add new institution..." option\n'
    '└── renderInstitutionTags()  → chips with × remove buttons\n'
    '\n'
    'handleInstitutionChange(sel)\n'
    '├── if value === "__add_new__" → prompt() for name\n'
    '├── add to state.institutions + saveState()\n'
    '└── populateInstitutionDropdown(newName)\n'
    '\n'
    'removeInstitution(inst)\n'
    '├── check not in use by any account\n'
    '└── filter out + saveState() + repopulate dropdown'
)

spacer()
heading('2.7 PIN / Lock Flow', 2)
code_block(
    'Modes: "setup" | "setup-confirm" | "unlock" | "change" | "change-confirm"\n'
    '\n'
    'First time (setup):\n'
    '  enter 4 digits → dots fill → 140ms delay → "setup-confirm"\n'
    '  confirm match  → SHA-256("pfmt:" + PIN) stored → _doUnlock()\n'
    '  confirm miss   → shake animation → back to "setup"\n'
    '\n'
    'Returning user (unlock):\n'
    '  enter PIN → hash compare → _doUnlock() → hide overlay\n'
    '  wrong PIN  → shake + "Wrong PIN. Try again."\n'
    '\n'
    'Change PIN:  "change" → verify current → "change-confirm" → update hash\n'
    'Forgot PIN:  localStorage.clear() → page reload'
)

spacer()
heading('2.8 ServiceNow Sync Flow', 2)
code_block(
    'testConnection()\n'
    '├── save credentials → state\n'
    '├── GET sys_user with Basic Auth\n'
    '├── try fetch PFMT tables to verify app installed\n'
    '└── setConnStatus("connected" | "error")\n'
    '\n'
    'syncFromSN()\n'
    '├── parallel fetch: x_1472763_person_0_transaction + x_1472763_person_0_account\n'
    '├── map SN fields → state object shapes\n'
    '└── renderAll() + showToast()\n'
    '\n'
    'pushAllToSN()\n'
    '├── filter local records (typeof id === "number")\n'
    '├── POST each to SN table\n'
    '├── replace local id with returned sys_id\n'
    '└── update state.lastSync'
)

spacer()
heading('2.9 AI (Groq) Call Flow', 2)
code_block(
    'callAI(prompt, targetElementId)\n'
    '├── check API key exists (warn if not)\n'
    '├── show 3-dot loading animation in targetElementId\n'
    '├── POST https://api.groq.com/openai/v1/chat/completions\n'
    '│   ├── model: llama-3.1-8b-instant\n'
    '│   └── max_tokens: 1000\n'
    '├── parse choices[0].message.content\n'
    '├── render markdown: **bold** → <strong>, \\n → <br>\n'
    '└── inject into targetElementId\n'
    '\n'
    'Entry points: openAI() | getBudgetAI() | getGoalsAI() | getAnalyticsAI() | getAccountsAI()'
)

spacer()
heading('2.10 Currency Switch Flow', 2)
code_block(
    'Currency select onchange:\n'
    '  state.currency = value → saveState() → renderAll()\n'
    '\n'
    'updateCurrencyDisplay()  [called inside renderAll()]\n'
    '├── currSym() → looks up CURRENCY_SYMBOLS[state.currency]\n'
    '├── update all .curr-sym elements with new symbol\n'
    '└── update all [data-curr-label] elements (replaces "CURR" with code)\n'
    '\n'
    'fmt(n) → currSym() + Math.abs(n).toLocaleString("en-SG", { 2 decimals })'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. DATA STRUCTURE
# ══════════════════════════════════════════════════════════════
heading('3. Data Structure', 1, C_JADE)

heading('3.1 State Object', 2)
code_block(
    'let state = {\n'
    '  transactions:  Transaction[],      // all transactions\n'
    '  budgets:       Budget[],           // monthly budget limits\n'
    '  goals:         Goal[],             // savings goals\n'
    '  accounts:      Account[],          // bank/wallet accounts\n'
    '  institutions:  string[],           // custom-managed institution names\n'
    '  currentType:   "expense"|"income"|"transfer",\n'
    '  nextId:        number,             // auto-increment counter\n'
    '  filterMonth:   "YYYY-MM",          // active month filter\n'
    '  geminiKey:     string,             // Groq API key\n'
    '  snInstance:    string,             // ServiceNow domain\n'
    '  snUsername:    string,\n'
    '  snPassword:    string,             // ⚠ stored plaintext\n'
    '  lastSync:      string | null,      // ISO timestamp\n'
    '  currency:      string,             // "SGD" | "USD" | ...\n'
    '  language:      "en" | "zh"\n'
    '};'
)

spacer()
heading('3.2 Object Shapes', 2)

body('Transaction', bold=True)
code_block(
    '{\n'
    '  id:          number,               // auto-incremented from state.nextId\n'
    '  type:        "expense"|"income"|"transfer",\n'
    '  amount:      number,               // always positive\n'
    '  description: string,               // e.g. "NTUC FairPrice"\n'
    '  category:    string,               // see categories list below\n'
    '  account:     string,               // matches Account.name\n'
    '  date:        "YYYY-MM-DD",\n'
    '  notes:       string                // optional\n'
    '}'
)

body('Budget', bold=True)
code_block(
    '{\n'
    '  id:        number,\n'
    '  category:  string,                 // unique per budget\n'
    '  amount:    number,                 // monthly limit\n'
    '  spent:     number,                 // calculated or manual override\n'
    '  alertPct:  number                  // alert threshold (default: 80)\n'
    '}'
)

body('Goal', bold=True)
code_block(
    '{\n'
    '  id:       number,\n'
    '  name:     string,                  // e.g. "Japan Trip"\n'
    '  icon:     string,                  // single emoji\n'
    '  target:   number,\n'
    '  current:  number,\n'
    '  monthly:  number,                  // suggested monthly contribution\n'
    '  date:     "YYYY-MM-DD"             // target completion date\n'
    '}'
)

body('Account', bold=True)
code_block(
    '{\n'
    '  id:          number,\n'
    '  name:        string,               // e.g. "DBS Checking"\n'
    '  type:        "Checking"|"Savings"|"Credit Card"|"Cash"|"Investment",\n'
    '  institution: string,               // from state.institutions[]\n'
    '  balance:     number,               // can be negative (debt)\n'
    '  icon:        "🏦" | "💳"\n'
    '}'
)

spacer()
heading('3.3 Categories', 2)
table_with_header(
    ['Type', 'Categories'],
    [
        ['Expense', 'Food & Drink, Transport, Groceries, Shopping, Bills, Health, Entertainment, Education, Travel, Other'],
        ['Income',  'Salary, Freelance, Investment, Other'],
    ]
)

spacer()
heading('3.4 localStorage Keys', 2)
table_with_header(
    ['Key', 'Contents'],
    [
        ['pfmt_state_v2', 'Full serialized state JSON (transactions, budgets, goals, accounts, institutions, settings)'],
        ['pfmt_pin',      'SHA-256 hash of "pfmt:" + 4-digit PIN'],
    ]
)

spacer()
heading('3.5 Currency Symbols Map', 2)
table_with_header(
    ['Currency', 'Symbol', 'Currency', 'Symbol'],
    [
        ['SGD', '$',    'EUR', '€'],
        ['USD', '$',    'GBP', '£'],
        ['AUD', 'A$',   'JPY', '¥'],
        ['MYR', 'RM',   'CNY', '¥'],
        ['HKD', 'HK$',  'IDR', 'Rp'],
        ['THB', '฿',    'PHP', '₱'],
        ['INR', '₹',    '',    ''],
    ]
)

spacer()
heading('3.6 Default Data', 2)
body('DEFAULT_INSTITUTIONS (line 1244):', bold=True)
bullet('DBS, OCBC, UOB, Standard Chartered, Grab Financial, Other')

body('DEFAULT_ACCOUNTS — loaded on first launch only:', bold=True)
table_with_header(
    ['Name', 'Type', 'Institution', 'Balance', 'Icon'],
    [
        ['DBS Checking',  'Checking',    'DBS',  '$4,820.50',  '🏦'],
        ['OCBC Savings',  'Savings',     'OCBC', '$12,500.00', '🏦'],
        ['Grab PayLater', 'Credit Card', 'Grab', '-$380.00',   '💳'],
    ]
)

spacer()
heading('3.7 TRANSLATIONS Structure', 2)
code_block(
    'const TRANSLATIONS = {\n'
    '  en: {\n'
    '    nav_dashboard:    "Dashboard",\n'
    '    lbl_net_balance:  "Net Balance",\n'
    '    toast_txn_saved:  "Transaction saved",\n'
    '    modal_add_txn:    "Add Transaction",\n'
    '    // ... 100+ keys\n'
    '  },\n'
    '  zh: {\n'
    '    nav_dashboard:    "仪表盘",\n'
    '    // ... parallel keys in Simplified Chinese\n'
    '  }\n'
    '};\n'
    '\n'
    '// Access:\n'
    'const t = key => (TRANSLATIONS[state.language])[key] || TRANSLATIONS.en[key] || key;'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. CODE REVIEW
# ══════════════════════════════════════════════════════════════
heading('4. Code Review & Architecture', 1, C_JADE)

heading('4.1 Key Functions Reference', 2)
table_with_header(
    ['Function', 'Purpose', 'Lines'],
    [
        ['renderAll()',              'Master render: applyLanguage + updateCurrencyDisplay + all 6 render*', '~1887'],
        ['renderDashboard()',        'KPIs, hero balance, recent txns, category chart, AI panel',            '~1900'],
        ['renderTransactions()',     'Filter by month/type/category/search, group by date',                  '~1966'],
        ['renderBudgets()',          'Budget cards with progress bars + monthly overview',                   '~1992'],
        ['renderGoals()',            'Goal cards with progress bars + summary panel',                        '~2052'],
        ['renderAnalytics()',        'All-time stats, largest expense, category breakdown',                  '~2112'],
        ['renderAccounts()',         'Account list, totals, insights panel',                                '~2138'],
        ['calcStats()',              'Income, expense, balance, savings rate, avg daily for active month',   '~1775'],
        ['getBudgetSpent(cat)',      'Sum expenses in current month for a category',                         '~1788'],
        ['fmt(n)',                   'currSym() + abs(n) formatted to 2 decimals',                          '~1725'],
        ['fmtDate(d)',               '"15 Jun 2026" from "2026-06-15"',                                     '~1735'],
        ['updateCurrencyDisplay()',  'Updates .curr-sym + [data-curr-label] elements',                       '~1727'],
        ['applyLanguage()',          'Updates all DOM text from TRANSLATIONS[state.language]',               '~1799'],
        ['openTxnModal(id?)',        'Opens Add/Edit transaction modal, populates fields',                   '~1363'],
        ['saveTransaction()',        'Validates, creates/updates txn, saves state',                         '~1436'],
        ['deleteTransaction(id)',    'Confirm dialog → remove → saveState → renderAll',                     '~1466'],
        ['showConfirm(title, msg)', 'Promise-based confirm dialog',                                         '~1327'],
        ['callAI(prompt, id)',       'POST to Groq API, render markdown response',                          '~2216'],
        ['buildSummary()',           'Context string from state data for AI prompts',                        '~2238'],
        ['loadSampleData()',         'Seeds demo data on first launch',                                     '~2460'],
        ['saveState()',              'Serialize state → localStorage[pfmt_state_v2]',                       '~1268'],
        ['loadState()',              'Deserialize + migrate from localStorage',                              '~1289'],
        ['testConnection()',         'Verify SN Basic Auth credentials',                                    '~2403'],
        ['syncFromSN()',             'Pull transactions + accounts from SN, map fields',                    '~2437'],
        ['pushAllToSN()',            'POST local records to SN, update ids to sys_ids',                     '~2288'],
        ['exportCSV()',              'Generate and auto-download transactions CSV',                          '~2247'],
        ['lockApp()',                'Show PIN unlock overlay',                                             '~2681'],
        ['_hashPin(pin)',            'SHA-256("pfmt:" + PIN)',                                              '~2544'],
    ]
)

spacer()
heading('4.2 Event Wiring Patterns', 2)
table_with_header(
    ['Pattern', 'Example'],
    [
        ['Navigation',           'onclick="showView(\'dashboard\', this)"'],
        ['Open modal',           'onclick="openTxnModal()"'],
        ['Edit with ID',         'onclick="openTxnModal(${txn.id})"'],
        ['Delete with confirm',  'onclick="deleteTransaction(${txn.id})"'],
        ['Filter inputs',        'oninput="renderTransactions()"'],
        ['Currency/Language',    'onchange="state.currency=this.value; saveState(); renderAll();"'],
        ['Institution select',   'onchange="handleInstitutionChange(this)"'],
        ['PIN numpad',           'onclick="pinInput(\'5\')"'],
        ['Keyboard shortcut',    'Ctrl+N → openTxnModal()  |  Escape → close all modals'],
        ['Backdrop click',       'Click outside modal → close'],
    ]
)

spacer()
heading('4.3 Issues & Recommendations', 2)
table_with_header(
    ['#', 'Issue', 'Severity', 'Recommendation'],
    [
        ['1', 'snPassword stored plaintext in localStorage',
              'High',
              'Use sessionStorage or prompt password on each session; never persist credentials'],
        ['2', 'No XSS sanitization on user input in HTML rendering (description, notes)',
              'Medium',
              'Use textContent or a sanitizer (DOMPurify) instead of innerHTML for user strings'],
        ['3', 'renderAll() re-renders all 6 views on every data change — no incremental updates',
              'Medium',
              'Only re-render the affected view; use dirty flags or targeted DOM updates'],
        ['4', 'Category names are magic strings duplicated in 4+ places',
              'Low',
              'Extract to a CATEGORIES constant object used everywhere'],
        ['5', 't() translation function shadowed by local "const t =" inside functions (was the edit/delete bug)',
              'Fixed',
              'Rename local transaction variables to "txn" (already done)'],
        ['6', 'No pagination on transaction list — could be slow with 1000+ records',
              'Low',
              'Add lazy-load or paginate at 50 records per page'],
        ['7', 'No rate limiting on Groq AI calls — user can spam the API key',
              'Low',
              'Add debounce + loading state lock to AI buttons'],
        ['8', 'Groq API key exposed in browser network requests (client-side)',
              'Medium',
              'Proxy through a backend or use a serverless function to hide the key'],
        ['9', 'Some error/toast strings not in TRANSLATIONS object (always English)',
              'Low',
              'Move all user-visible strings into TRANSLATIONS for full i18n'],
        ['10', 'Missing ARIA labels on icon buttons and PIN dots',
               'Low',
               'Add aria-label to all icon-only buttons; announce PIN status to screen readers'],
    ],
    header_bg='FF5B5B'
)

spacer()
heading('4.4 Hardcoded Values', 2)
table_with_header(
    ['Value', 'Location', 'Notes'],
    [
        ['dev275144.service-now.com',         'state default + multiple links', 'Default SN PDI instance — should be user-configurable only'],
        ['x_1472763_person_0_transaction',    'SN sync functions',             'SN table name — consider making configurable'],
        ['x_1472763_person_0_account',        'SN sync functions',             'SN table name'],
        ['https://api.groq.com/openai/v1/...','callAI()',                      'Groq endpoint — ok to hardcode'],
        ['llama-3.1-8b-instant',              'callAI()',                      'Groq model — consider exposing in Settings'],
        ['pfmt_state_v2',                     'saveState()/loadState()',        'localStorage key'],
        ['pfmt_pin',                          'PIN functions',                 'localStorage key for PIN hash'],
        ['4',                                 'PIN validation',                'PIN length (line ~2560)'],
    ]
)

spacer()
heading('4.5 Architecture Summary', 2)
body('Strengths:', bold=True)
bullet('Single-file HTML/CSS/JS — zero build tooling, instant deploy to GitHub Pages')
bullet('Centralized state object — predictable data flow')
bullet('Translation layer covers English and Simplified Chinese')
bullet('Responsive mobile-first design with touch-friendly adjustments')
bullet('ServiceNow integration for enterprise data sync')
bullet('SHA-256 PIN hashing for local security')
bullet('Dynamic currency system with 13 supported currencies')

spacer()
body('Limitations to address for production:', bold=True)
bullet('No backend — all data is local; no multi-device sync without SN')
bullet('No user auth — single PIN protects entire app on device')
bullet('No real-time updates — all renders are triggered manually')
bullet('No offline-first strategy — SN sync fails silently offline')
bullet('Plaintext password storage in localStorage is a security risk')

spacer()
spacer()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run('Personal Money Tracker  ·  Technical Documentation  ·  June 2026')
run_f.font.size = Pt(8)
run_f.font.color.rgb = C_MUTED

# ── Save ──────────────────────────────────────────────────────
out = '/home/user/Personal-Money-Tracker/Personal_Money_Tracker_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
